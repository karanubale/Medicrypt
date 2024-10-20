import os
import requests
from flask import Flask, request, jsonify
import speech_recognition as sr
from docx import Document
import logging
from logging.handlers import RotatingFileHandler
from dotenv import load_dotenv
from datetime import datetime
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables from a .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Load configuration from environment variables
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1")

# Prompt for the AI assistant
prompt = '''
You are an AI medical assistant designed to generate concise prescription reports from doctor-patient conversation transcripts. Your task is to extract only the medically relevant information needed for the prescription, ignoring any tangents or small talk.

**Instructions:**
Given the conversation transcript, summarize and organize the following information into a prescription report:

### **Prescription Report:**
- **Patient's Name:**
- **Patient's Age:**
- **Patient's Gender:**
- **Chief Complaint/Reason for Visit:**
- **Relevant Medical History:**
- **Current Medications:**
- **Allergies:**
- **Vital Signs (if provided):**
- **Physical Exam Findings (if provided):**
- **Diagnosis:**
- **Treatment Plan:** (Include drug name, dosage, frequency, and duration)
- **Follow-up Instructions:**

**Notes:**

1. Extract only the relevant details and Exclude any irrelevant conversation.
2. Return only the report itself without additional commentary.
'''

# Configure logging
logging.basicConfig(level=logging.INFO)
handler = RotatingFileHandler("app.log", maxBytes=100000, backupCount=3)
handler.setLevel(logging.INFO)
app.logger.addHandler(handler)

def create_document_from_groq_response(response_text, document_name):
    document = Document()
    
    # Add Title with formatting
    title = document.add_heading('Prescription Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add Date and Time
    now = datetime.now()
    date_time_str = now.strftime("%Y-%m-%d | %H:%M:%S")
    document.add_paragraph(f"Date and Time: {date_time_str}")
    
    # Parse and format response text
    response_lines = response_text.split("\n")
    for line in response_lines:
        if line.startswith("- **"):
            # Add each prescription item as a separate paragraph with some space before it
            para = document.add_paragraph(line)
            para.paragraph_format.space_before = Pt(12)
        else:
            document.add_paragraph(line)

    # Save the document
    document_path = os.path.join(os.getcwd(), f"{document_name}.docx")
    document.save(document_path)
    return document_path

def get_groq_response(text):
    """Fetch the Groq API response for the given text."""
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama3-70b-8192",
        "messages": [
            {
                "role": "system",
                "content": prompt,
            },
            {
                "role": "user",
                "content": text
            }
        ]
    }
    try:
        response = requests.post(f"{GROQ_BASE_URL}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except requests.exceptions.RequestException as e:
        app.logger.error(f"Error calling Groq API: {e}")
        return str(e)

@app.route('/process_audio', methods=['POST'])
def process_audio():
    """Endpoint to process an audio file and return a prescription report."""
    if 'audio' not in request.files:
        app.logger.error('No audio file part in the request')
        return jsonify({'error': 'No audio file part'}), 400
    
    file = request.files['audio']
    if file.filename == '':
        app.logger.error('No selected file')
        return jsonify({'error': 'No selected file'}), 400

    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(file) as source:
            audio = recognizer.listen(source)
        text = recognizer.recognize_google(audio)
        groq_response = get_groq_response(text)
        document_name = "Groq_Prescription_Report"
        document_path = create_document_from_groq_response(groq_response, document_name)
        return jsonify({'prescription': groq_response, 'document_path': document_path})
    except sr.UnknownValueError:
        app.logger.error('Audio unintelligible')
        return jsonify({'error': 'Audio unintelligible'}), 400
    except sr.RequestError as e:
        app.logger.error(f"Speech recognition error: {e}")
        return jsonify({'error': 'Speech recognition service error'}), 500
    except Exception as e:
        app.logger.error(f"Unexpected error: {e}")
        return jsonify({'error': str(e)}), 500



@app.route('/process_text', methods=['POST'])
def process_text():
    """Endpoint to process a text input and return a prescription report."""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400

    try:
        text = data['text']
        groq_response = get_groq_response(text)
        document_name = "Groq_Prescription_Report"
        document_path = create_document_from_groq_response(groq_response, document_name)
        return jsonify({'prescription': groq_response, 'document_path': document_path})
    except Exception as e:
        app.logger.error(f"Unexpected error: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv("PORT", 5000)), debug=os.getenv("DEBUG", "False") == "True")
